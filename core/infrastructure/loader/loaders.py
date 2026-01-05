"""
Format-specific Document Loaders
格式特定的文件載入器

Implements loaders for various file formats:
- TXT: Plain text
- CSV: Comma-separated values
- XLSX/XLS: Excel spreadsheets
- DOCX: Microsoft Word
- JSON: Structured JSON
- PDF: Portable Document Format
- HTML: Web documents
- XML: Structured XML
- FHIR: Healthcare data standard
"""

import csv
import json
from pathlib import Path
from typing import Any

from loguru import logger

from .base import (
    BaseBinaryLoader,
    BaseTextLoader,
    DocumentFormat,
    DocumentMetadata,
    LoadedDocument,
)


class TextLoader(BaseTextLoader):
    """
    Plain text file loader
    純文字檔案載入器
    
    Loads .txt files with automatic encoding detection.
    
    Examples:
        >>> loader = TextLoader()
        >>> doc = loader.load("patient_record.txt")
        >>> print(doc.content)
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load TXT file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading TXT file: {file_path}")

        # Read content
        content = self._read_text_file(file_path)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(content=content, metadata=metadata)

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports TXT format"""
        return format == DocumentFormat.TXT


class CSVLoader(BaseTextLoader):
    """
    CSV file loader
    CSV 檔案載入器
    
    Loads CSV files and converts to structured data.
    Can return either:
    - Single concatenated text (all rows joined)
    - Multiple records (list of dictionaries)
    
    Examples:
        >>> loader = CSVLoader()
        >>> doc = loader.load("patient_list.csv")
        >>> print(doc.records[0])  # First patient record
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load CSV file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading CSV file: {file_path}")

        # Read CSV
        with open(file_path, encoding=self.config.encoding,
                 errors=self.config.encoding_errors) as f:

            csv_reader = csv.DictReader(
                f,
                delimiter=self.config.csv_delimiter,
                quotechar=self.config.csv_quotechar
            ) if self.config.csv_has_header else csv.reader(
                f,
                delimiter=self.config.csv_delimiter,
                quotechar=self.config.csv_quotechar
            )

            records = list(csv_reader)

        # Convert to text (concatenate all rows)
        if self.config.csv_has_header and records:
            # Format: "field1: value1, field2: value2\n"
            content_lines = []
            for record in records:
                line = ", ".join([f"{k}: {v}" for k, v in record.items()])
                content_lines.append(line)
            content = "\n".join(content_lines)
        else:
            # Format: "value1, value2, value3\n"
            content = "\n".join([", ".join(row) for row in records])

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata_dict["custom"] = {
            "num_records": len(records),
            "delimiter": self.config.csv_delimiter
        }
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(
            content=content,
            metadata=metadata,
            records=records if self.config.csv_has_header else None
        )

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports CSV format"""
        return format == DocumentFormat.CSV


class ExcelLoader(BaseBinaryLoader):
    """
    Excel file loader (XLSX/XLS)
    Excel 檔案載入器
    
    Loads Excel files with multi-sheet support.
    Uses openpyxl for .xlsx and xlrd for .xls
    
    Examples:
        >>> loader = ExcelLoader()
        >>> # Load all sheets
        >>> doc = loader.load("patient_data.xlsx")
        >>> 
        >>> # Load specific sheet
        >>> config = LoaderConfig(excel_sheet_name="Patient Info")
        >>> loader = ExcelLoader(config)
        >>> doc = loader.load("patient_data.xlsx")
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load Excel file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading Excel file: {file_path}")

        # Determine Excel type
        if file_path.suffix.lower() == '.xlsx':
            return self._load_xlsx(file_path)
        elif file_path.suffix.lower() == '.xls':
            return self._load_xls(file_path)
        else:
            raise ValueError(f"Unsupported Excel format: {file_path.suffix}")

    def _load_xlsx(self, file_path: Path) -> LoadedDocument:
        """Load XLSX file using openpyxl"""
        try:
            import openpyxl
        except ImportError:
            raise ImportError(
                "openpyxl is required for XLSX files. "
                "Install with: pip install openpyxl"
            )

        workbook = openpyxl.load_workbook(file_path, data_only=True)

        # Get sheets to load
        if self.config.excel_sheet_name:
            if self.config.excel_sheet_name not in workbook.sheetnames:
                raise ValueError(
                    f"Sheet '{self.config.excel_sheet_name}' not found. "
                    f"Available: {workbook.sheetnames}"
                )
            sheets_to_load = [self.config.excel_sheet_name]
        else:
            sheets_to_load = workbook.sheetnames

        # Extract data from sheets
        all_content = []
        all_records = []

        for sheet_name in sheets_to_load:
            sheet = workbook[sheet_name]

            # Skip rows if configured
            rows = list(sheet.iter_rows(
                min_row=self.config.excel_skip_rows + 1,
                values_only=True
            ))

            if not rows:
                continue

            # Get header (first row after skip)
            header = [str(cell) if cell is not None else f"Column_{i}"
                     for i, cell in enumerate(rows[0])]

            # Process data rows
            for row in rows[1:]:
                # Create record
                record = {}
                for i, value in enumerate(row):
                    if i < len(header):
                        record[header[i]] = str(value) if value is not None else ""

                all_records.append(record)

                # Create text representation
                row_text = ", ".join([f"{k}: {v}" for k, v in record.items()])
                all_content.append(f"[{sheet_name}] {row_text}")

        content = "\n".join(all_content)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata_dict["custom"] = {
            "num_sheets": len(sheets_to_load),
            "sheet_names": sheets_to_load,
            "num_records": len(all_records)
        }
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(
            content=content,
            metadata=metadata,
            records=all_records
        )

    def _load_xls(self, file_path: Path) -> LoadedDocument:
        """Load XLS file using xlrd"""
        try:
            import xlrd
        except ImportError:
            raise ImportError(
                "xlrd is required for XLS files. "
                "Install with: pip install xlrd"
            )

        workbook = xlrd.open_workbook(file_path)

        # Get sheets to load
        if self.config.excel_sheet_name:
            try:
                sheet_index = workbook.sheet_names().index(self.config.excel_sheet_name)
                sheets_to_load = [workbook.sheet_by_index(sheet_index)]
            except ValueError:
                raise ValueError(
                    f"Sheet '{self.config.excel_sheet_name}' not found. "
                    f"Available: {workbook.sheet_names()}"
                )
        else:
            sheets_to_load = [workbook.sheet_by_index(i)
                            for i in range(workbook.nsheets)]

        # Extract data
        all_content = []
        all_records = []

        for sheet in sheets_to_load:
            if sheet.nrows <= self.config.excel_skip_rows:
                continue

            # Get header
            header_row = self.config.excel_skip_rows
            header = [str(sheet.cell_value(header_row, col))
                     for col in range(sheet.ncols)]

            # Process data rows
            for row_idx in range(header_row + 1, sheet.nrows):
                record = {}
                for col_idx in range(sheet.ncols):
                    value = sheet.cell_value(row_idx, col_idx)
                    record[header[col_idx]] = str(value) if value else ""

                all_records.append(record)

                row_text = ", ".join([f"{k}: {v}" for k, v in record.items()])
                all_content.append(f"[{sheet.name}] {row_text}")

        content = "\n".join(all_content)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata_dict["custom"] = {
            "num_sheets": len(sheets_to_load),
            "num_records": len(all_records)
        }
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(
            content=content,
            metadata=metadata,
            records=all_records
        )

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports XLSX and XLS formats"""
        return format in [DocumentFormat.XLSX, DocumentFormat.XLS]


class WordLoader(BaseBinaryLoader):
    """
    Microsoft Word document loader (DOCX)
    Word 文件載入器
    
    Loads .docx files and extracts text content.
    Uses python-docx library.
    
    Examples:
        >>> loader = WordLoader()
        >>> doc = loader.load("medical_report.docx")
        >>> print(doc.content)
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load DOCX file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading Word document: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX files. "
                "Install with: pip install python-docx"
            )

        # Load document
        doc = Document(file_path)

        # Extract paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extract tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                table_texts.append(row_text)

        # Combine content
        content_parts = paragraphs
        if table_texts:
            content_parts.append("\n--- Tables ---\n")
            content_parts.extend(table_texts)

        content = "\n".join(content_parts)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        core_properties = doc.core_properties
        metadata_dict["custom"] = {
            "num_paragraphs": len(paragraphs),
            "num_tables": len(doc.tables),
            "author": core_properties.author if hasattr(core_properties, 'author') else None,
            "title": core_properties.title if hasattr(core_properties, 'title') else None,
        }
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(content=content, metadata=metadata)

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports DOCX format"""
        return format == DocumentFormat.DOCX


class JSONLoader(BaseTextLoader):
    """
    JSON file loader
    JSON 檔案載入器
    
    Loads JSON files and converts to both text and structured data.
    
    Examples:
        >>> loader = JSONLoader()
        >>> doc = loader.load("patient.json")
        >>> print(doc.structured_data["patient_id"])
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load JSON file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading JSON file: {file_path}")

        # Read JSON
        with open(file_path, encoding=self.config.encoding) as f:
            data = json.load(f)

        # Convert to text (pretty JSON)
        content = json.dumps(data, ensure_ascii=False, indent=2)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(
            content=content,
            metadata=metadata,
            structured_data=data
        )

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports JSON format"""
        return format == DocumentFormat.JSON


class PDFLoader(BaseBinaryLoader):
    """
    PDF file loader
    PDF 檔案載入器
    
    Loads PDF files and extracts text content.
    Uses PyPDF2 or pdfplumber.
    
    Examples:
        >>> loader = PDFLoader()
        >>> doc = loader.load("lab_results.pdf")
        >>> print(doc.content)
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load PDF file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading PDF file: {file_path}")

        try:
            import pdfplumber
            return self._load_with_pdfplumber(file_path)
        except ImportError:
            try:
                import PyPDF2
                return self._load_with_pypdf2(file_path)
            except ImportError:
                raise ImportError(
                    "PDF loading requires pdfplumber or PyPDF2. "
                    "Install with: pip install pdfplumber or pip install PyPDF2"
                )

    def _load_with_pdfplumber(self, file_path: Path) -> LoadedDocument:
        """Load PDF using pdfplumber (better text extraction)"""
        import pdfplumber

        all_text = []
        num_pages = 0

        with pdfplumber.open(file_path) as pdf:
            pages_to_extract = range(len(pdf.pages))

            if self.config.pdf_page_range:
                start, end = self.config.pdf_page_range
                pages_to_extract = range(start - 1, min(end, len(pdf.pages)))

            for page_num in pages_to_extract:
                page = pdf.pages[page_num]
                text = page.extract_text()
                if text:
                    all_text.append(f"--- Page {page_num + 1} ---\n{text}")
                num_pages += 1

        content = "\n\n".join(all_text)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata_dict["custom"] = {"num_pages": num_pages}
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(content=content, metadata=metadata)

    def _load_with_pypdf2(self, file_path: Path) -> LoadedDocument:
        """Load PDF using PyPDF2 (fallback)"""
        import PyPDF2

        all_text = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            pages_to_extract = range(num_pages)
            if self.config.pdf_page_range:
                start, end = self.config.pdf_page_range
                pages_to_extract = range(start - 1, min(end, num_pages))

            for page_num in pages_to_extract:
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    all_text.append(f"--- Page {page_num + 1} ---\n{text}")

        content = "\n\n".join(all_text)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata_dict["custom"] = {"num_pages": len(pages_to_extract)}
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(content=content, metadata=metadata)

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports PDF format"""
        return format == DocumentFormat.PDF


class HTMLLoader(BaseTextLoader):
    """
    HTML file loader
    HTML 檔案載入器
    
    Loads HTML files and extracts text content.
    Uses BeautifulSoup for parsing.
    
    Examples:
        >>> loader = HTMLLoader()
        >>> doc = loader.load("patient_portal.html")
        >>> print(doc.content)
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load HTML file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading HTML file: {file_path}")

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "BeautifulSoup is required for HTML files. "
                "Install with: pip install beautifulsoup4"
            )

        # Read HTML
        html_content = self._read_text_file(file_path)

        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove scripts and styles if configured
        if self.config.html_remove_scripts:
            for script in soup(["script", "style"]):
                script.decompose()

        # Extract text
        if self.config.html_extract_text_only:
            content = soup.get_text(separator='\n', strip=True)
        else:
            content = str(soup)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)

        # Try to extract HTML metadata
        title = soup.find('title')
        metadata_dict["custom"] = {
            "title": title.string if title else None,
        }

        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(content=content, metadata=metadata)

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports HTML format"""
        return format == DocumentFormat.HTML


class XMLLoader(BaseTextLoader):
    """
    XML file loader
    XML 檔案載入器
    
    Loads XML files and converts to text representation.
    Uses xml.etree.ElementTree for parsing.
    
    Examples:
        >>> loader = XMLLoader()
        >>> doc = loader.load("lab_data.xml")
        >>> print(doc.structured_data)
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load XML file"""
        file_path = Path(file_path)
        self._validate_file(file_path)

        logger.info(f"Loading XML file: {file_path}")

        import xml.etree.ElementTree as ET

        # Parse XML
        tree = ET.parse(file_path)
        root = tree.getroot()

        # Convert to text representation
        content_lines = []
        self._xml_to_text(root, content_lines, indent=0)
        content = "\n".join(content_lines)

        # Convert to dictionary
        structured_data = self._xml_to_dict(root)

        # Extract metadata
        metadata_dict = self._extract_file_metadata(file_path)
        metadata_dict["custom"] = {
            "root_tag": root.tag,
            "num_children": len(root)
        }
        metadata = DocumentMetadata(**metadata_dict)

        return LoadedDocument(
            content=content,
            metadata=metadata,
            structured_data=structured_data
        )

    def _xml_to_text(self, element, lines: list[str], indent: int = 0):
        """Convert XML element to text lines"""
        prefix = "  " * indent

        # Element tag and attributes
        if element.attrib:
            attr_str = " ".join([f'{k}="{v}"' for k, v in element.attrib.items()])
            lines.append(f"{prefix}{element.tag} [{attr_str}]")
        else:
            lines.append(f"{prefix}{element.tag}")

        # Element text
        if element.text and element.text.strip():
            lines.append(f"{prefix}  {element.text.strip()}")

        # Children
        for child in element:
            self._xml_to_text(child, lines, indent + 1)

    def _xml_to_dict(self, element) -> dict[str, Any]:
        """Convert XML element to dictionary"""
        result = {}

        # Attributes
        if element.attrib:
            result["@attributes"] = dict(element.attrib)

        # Text content
        if element.text and element.text.strip():
            result["#text"] = element.text.strip()

        # Children
        for child in element:
            child_dict = self._xml_to_dict(child)

            if child.tag in result:
                # Multiple children with same tag -> list
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_dict)
            else:
                result[child.tag] = child_dict

        return {element.tag: result} if result else {element.tag: element.text}

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports XML format"""
        return format == DocumentFormat.XML


class FHIRLoader(JSONLoader):
    """
    FHIR (Fast Healthcare Interoperability Resources) loader
    FHIR 醫療資料載入器
    
    Loads FHIR R4 JSON resources.
    Inherits from JSONLoader but adds FHIR-specific validation.
    
    Examples:
        >>> loader = FHIRLoader()
        >>> doc = loader.load("patient_fhir.json")
        >>> print(doc.structured_data["resourceType"])
    """

    def load(self, file_path: str | Path) -> LoadedDocument:
        """Load FHIR JSON file"""
        # Use parent JSONLoader
        doc = super().load(file_path)

        # Validate FHIR structure
        if doc.structured_data:
            if "resourceType" not in doc.structured_data:
                logger.warning(f"FHIR resource missing 'resourceType': {file_path}")
            else:
                resource_type = doc.structured_data["resourceType"]
                logger.info(f"Loaded FHIR resource type: {resource_type}")

                # Update metadata
                doc.metadata.custom["fhir_resource_type"] = resource_type
                if "id" in doc.structured_data:
                    doc.metadata.custom["fhir_resource_id"] = doc.structured_data["id"]

        return doc

    def supports_format(self, format: DocumentFormat) -> bool:
        """Supports FHIR format (JSON-based)"""
        return format == DocumentFormat.FHIR
